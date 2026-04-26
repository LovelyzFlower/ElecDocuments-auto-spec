import sys
import os
import subprocess
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                               QHBoxLayout, QPushButton, QLabel, QFileDialog, 
                               QTableWidget, QTableWidgetItem, QHeaderView, 
                               QComboBox, QMessageBox, QSplitter, QGraphicsView,
                               QGraphicsScene, QGraphicsPixmapItem, QGraphicsRectItem,
                               QProgressBar, QCheckBox)
from PySide6.QtGui import QPixmap, QImage, QPen, QColor, QBrush
from PySide6.QtCore import Qt, QThread, Signal, QRectF, QTimer

from ocr_engine import OCREngine
from matcher import SemanticMatcher
from utils import load_metadata, save_spec, draw_bboxes_on_image
import fitz
from docx import Document

class WorkerThread(QThread):
    finished = Signal(object)
    error = Signal(str)

    def __init__(self, func, *args, **kwargs):
        super().__init__()
        self.func = func
        self.args = args
        self.kwargs = kwargs

    def run(self):
        try:
            result = self.func(*self.args, **self.kwargs)
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))

class InteractiveRectItem(QGraphicsRectItem):
    def __init__(self, rect, row_idx, callback):
        super().__init__(rect)
        self.row_idx = row_idx
        self.callback = callback
        self.is_included = True
        
        self.setFlag(QGraphicsRectItem.ItemIsSelectable, True)
        self.update_style()
        
    def update_style(self):
        if self.is_included:
            self.setPen(QPen(QColor(0, 255, 0), 3))
            self.setBrush(QBrush(QColor(0, 255, 0, 50)))
        else:
            self.setPen(QPen(QColor(255, 0, 0), 3))
            self.setBrush(QBrush(QColor(255, 0, 0, 50)))
            
    def mousePressEvent(self, event):
        self.is_included = not self.is_included
        self.update_style()
        self.callback(self.row_idx, self.is_included)
        super().mousePressEvent(event)

    def blink(self):
        self.blink_count = 0
        if hasattr(self, 'blink_timer'):
            self.blink_timer.stop()
        self.blink_timer = QTimer()
        self.blink_timer.timeout.connect(self._toggle_blink)
        self.blink_timer.start(150)

    def _toggle_blink(self):
        self.blink_count += 1
        if self.blink_count % 2 == 1:
            self.setPen(QPen(QColor(255, 255, 0), 5)) # Yellow
            self.setBrush(QBrush(QColor(255, 255, 0, 150)))
        else:
            self.update_style()
            
        if self.blink_count >= 6: # 3 flashes
            self.blink_timer.stop()

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Auto-Spec: AI-based E-Form Variable Specifier")
        self.resize(1200, 900) # Increased size for better A4 viewing

        # Main Layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        # Splitter for left (image) and right (table)
        splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(splitter)

        # Left Panel (Image Viewer)
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        
        self.graphics_view = QGraphicsView()
        self.graphics_scene = QGraphicsScene()
        self.graphics_view.setScene(self.graphics_scene)
        self.graphics_view.setStyleSheet("border: 1px solid #ccc; background-color: #f9f9f9;")
        left_layout.addWidget(self.graphics_view)
        
        self.rect_items = []
        
        splitter.addWidget(left_panel)

        # Right Panel (Data Table)
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        
        self.table_widget = QTableWidget()
        self.table_widget.setColumnCount(6)
        self.table_widget.setHorizontalHeaderLabels(["포함", "인식된 텍스트", "추천 변수명(한글)", "추천 변수명(영문)", "유사도", "수동 선택"])
        self.table_widget.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table_widget.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        
        # Connect selection change to blink functionality
        self.table_widget.itemSelectionChanged.connect(self.on_table_selection_changed)
        
        right_layout.addWidget(self.table_widget)
        splitter.addWidget(right_panel)

        # Set Splitter ratio to strongly vertical proportion (e.g. 450 width vs 750 width for right)
        splitter.setSizes([450, 750])

        # Bottom Panel (Controls)
        bottom_panel = QWidget()
        bottom_layout = QHBoxLayout(bottom_panel)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)
        main_layout.addWidget(bottom_panel)

        self.btn_load_meta = QPushButton("1. 메타데이터 로드")
        self.btn_load_image = QPushButton("2. 서식 문서 로드")
        self.btn_extract = QPushButton("3. 텍스트 추출 (OCR)")
        self.btn_map = QPushButton("4. 매핑 실행")
        self.btn_save = QPushButton("5. 명세서 저장")

        bottom_layout.addWidget(self.btn_load_meta)
        bottom_layout.addWidget(self.btn_load_image)
        bottom_layout.addWidget(self.btn_extract)
        bottom_layout.addWidget(self.btn_map)
        bottom_layout.addStretch()
        bottom_layout.addWidget(self.btn_save)

        # Connect signals
        self.btn_load_meta.clicked.connect(self.load_metadata_action)
        self.btn_load_image.clicked.connect(self.load_document_action)
        self.btn_extract.clicked.connect(self.run_extraction_action)
        self.btn_map.clicked.connect(self.run_map_action)
        self.btn_save.clicked.connect(self.save_spec_action)

        # Initial state
        self.btn_extract.setEnabled(False)
        self.btn_map.setEnabled(False)
        self.btn_save.setEnabled(False)

        # Application State
        self.metadata_df = None
        self.metadata_column = None # the column to match against
        self.metadata_mapping = {} # Korean -> English mapping
        self.image_path = None
        self.pdf_page_info = [] # List of (image_path, height, y_offset)
        self.ocr_results = []

        # Engines (Lazy loading)
        self.ocr_engine = None
        self.matcher = None

    def load_metadata_action(self):
        file_name, _ = QFileDialog.getOpenFileName(self, "메타데이터 열기", "", "Excel Files (*.xlsx *.xls);;JSON Files (*.json);;CSV Files (*.csv)")
        if file_name:
            try:
                self.metadata_df = load_metadata(file_name)
                # First column is English, second column is Korean
                eng_col = self.metadata_df.columns[0]
                kor_col = self.metadata_df.columns[1] if len(self.metadata_df.columns) > 1 else self.metadata_df.columns[0]
                self.metadata_column = kor_col
                
                self.metadata_mapping = {}
                for _, row in self.metadata_df.iterrows():
                    kor = str(row[kor_col])
                    eng = str(row[eng_col])
                    if kor and kor != 'nan':
                        self.metadata_mapping[kor] = eng if eng != 'nan' else ""
                        
                metadata_items = list(self.metadata_mapping.keys())
                
                # Initialize matcher if not already
                if self.matcher is None:
                    self.statusBar().showMessage("Semantic Matcher 모델 로딩 중... (처음 실행 시 다소 시간이 걸릴 수 있습니다.)")
                    self.progress_bar.setVisible(True)
                    self.progress_bar.setRange(0, 0)
                    QApplication.processEvents()
                    self.matcher = SemanticMatcher()
                    self.progress_bar.setVisible(False)

                self.matcher.fit_metadata(metadata_items)
                
                QMessageBox.information(self, "성공", f"메타데이터 로드 완료! (항목 수: {len(metadata_items)})")
                self.statusBar().showMessage("메타데이터 로드 완료.")
                self.check_ready_state()
            except Exception as e:
                QMessageBox.critical(self, "오류", f"메타데이터 로드 실패: {e}")

    def load_document_action(self):
        file_name, _ = QFileDialog.getOpenFileName(self, "이미지/문서 열기", "", "Supported Files (*.png *.jpg *.jpeg *.bmp *.pdf *.docx);;Image Files (*.png *.jpg *.jpeg *.bmp);;PDF Files (*.pdf);;DOCX Files (*.docx)")
        if file_name:
            self.image_path = file_name
            self.graphics_scene.clear()
            self.rect_items = []
            self.pdf_page_info = []
            
            if file_name.lower().endswith('.pdf'):
                self.statusBar().showMessage("PDF 로드 중...")
                self.progress_bar.setVisible(True)
                self.progress_bar.setRange(0, 0)
                QApplication.processEvents()
                # Render all pages of PDF
                try:
                    doc = fitz.open(file_name)
                    y_offset = 0
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap(dpi=150)
                        temp_img_path = os.path.join(os.path.dirname(file_name), f"temp_pdf_page_{page_num}.png")
                        pix.save(temp_img_path)
                        
                        qpixmap = QPixmap(temp_img_path)
                        pix_item = self.graphics_scene.addPixmap(qpixmap)
                        pix_item.setPos(0, y_offset)
                        
                        self.pdf_page_info.append((temp_img_path, pix.height, y_offset))
                        y_offset += pix.height + 20 # 20px padding between pages
                        
                    self.statusBar().showMessage(f"PDF 로드 완료 ({len(doc)}페이지).")
                    self.fit_to_width()
                except Exception as e:
                    QMessageBox.critical(self, "오류", f"PDF 로드 실패: {e}")
            elif file_name.lower().endswith('.docx'):
                self.graphics_scene.addText("DOCX 문서가 로드되었습니다.\n(미리보기는 지원되지 않습니다.)")
                self.statusBar().showMessage("DOCX 문서 로드 완료.")
            else:
                pixmap = QPixmap(self.image_path)
                self.graphics_scene.addPixmap(pixmap)
                self.fit_to_width()
                self.statusBar().showMessage("이미지 로드 완료.")
                
            self.progress_bar.setVisible(False)
            self.check_ready_state()

    def resizeEvent(self, event):
        self.fit_to_width()
        super().resizeEvent(event)

    def fit_to_width(self):
        if not self.graphics_scene.items():
            return
        rect = self.graphics_scene.itemsBoundingRect()
        if rect.width() == 0:
            return
        
        # Calculate scale factor to fit the width. 
        # Subtracting 5 pixels to prevent horizontal scrollbar.
        view_width = self.graphics_view.viewport().width() - 5
        scale_factor = view_width / rect.width()
        
        self.graphics_view.resetTransform()
        self.graphics_view.scale(scale_factor, scale_factor)
        self.graphics_view.setAlignment(Qt.AlignTop | Qt.AlignHCenter)

    def check_ready_state(self):
        if self.metadata_df is not None and self.image_path is not None:
            self.btn_extract.setEnabled(True)

    def run_extraction_action(self):
        if self.ocr_engine is None:
            self.statusBar().showMessage("OCR 모델 로딩 중...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)
            QApplication.processEvents()
            self.ocr_engine = OCREngine()
            
        self.statusBar().showMessage("텍스트 추출 중...")
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)
        self.btn_extract.setEnabled(False)
        self.btn_map.setEnabled(False)
        self.btn_load_meta.setEnabled(False)
        self.btn_load_image.setEnabled(False)
        
        self.worker = WorkerThread(self._extract_text)
        self.worker.finished.connect(self._on_extraction_finished)
        self.worker.error.connect(self._on_processing_error)
        self.worker.start()

    def _extract_text(self):
        if self.image_path.lower().endswith('.docx'):
            doc = Document(self.image_path)
            extracted_texts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    extracted_texts.append(text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and text not in extracted_texts:
                            extracted_texts.append(text)
                            
            # Create pseudo OCR results
            ocr_res = [([[0,0], [0,0], [0,0], [0,0]], t, 1.0) for t in extracted_texts]
        elif self.image_path.lower().endswith('.pdf'):
            all_ocr_res = []
            for path, height, y_offset in self.pdf_page_info:
                page_res = self.ocr_engine.extract_text(path)
                for item in page_res:
                    bbox, text, prob = item
                    # Adjust y coordinates based on the page's vertical offset
                    adjusted_bbox = [[x, y + y_offset] for x, y in bbox]
                    all_ocr_res.append((adjusted_bbox, text, prob))
            ocr_res = all_ocr_res
        else:
            # 1. Single Image OCR Extraction
            ocr_res = self.ocr_engine.extract_text(self.image_path)
            
        return ocr_res

    def _on_extraction_finished(self, ocr_results):
        self.ocr_results = ocr_results
        self.rect_items = []
        
        # Draw bboxes on scene
        if self.graphics_scene.items() and not self.image_path.lower().endswith('.docx'):
            # remove old rects
            for item in self.graphics_scene.items():
                if isinstance(item, InteractiveRectItem):
                    self.graphics_scene.removeItem(item)
                    
            for row_idx, item in enumerate(self.ocr_results):
                bbox, text, prob = item
                x1, y1 = bbox[0][0], bbox[0][1]
                x2, y2 = bbox[2][0], bbox[2][1]
                rect = QRectF(x1, y1, x2 - x1, y2 - y1)
                rect_item = InteractiveRectItem(rect, row_idx, self.on_rect_toggled)
                self.graphics_scene.addItem(rect_item)
                self.rect_items.append(rect_item)

        # Populate Table
        self.table_widget.setRowCount(len(self.ocr_results))
        for row_idx, item in enumerate(self.ocr_results):
            text = item[1]
            
            checkbox = QCheckBox()
            checkbox.setChecked(True)
            checkbox.setStyleSheet("margin-left: 5px; margin-right: 5px;")
            checkbox.stateChanged.connect(lambda state, r=row_idx: self.on_table_checkbox_toggled(r, state))
            
            cell_widget = QWidget()
            layout = QHBoxLayout(cell_widget)
            layout.addWidget(checkbox)
            layout.setAlignment(Qt.AlignCenter)
            layout.setContentsMargins(0,0,0,0)
            self.table_widget.setCellWidget(row_idx, 0, cell_widget)
            
            text_item = QTableWidgetItem(text)
            text_item.setFlags(text_item.flags() & ~Qt.ItemIsEditable)
            self.table_widget.setItem(row_idx, 1, text_item)
            
            # Clear matching columns
            for col in range(2, 6):
                self.table_widget.setItem(row_idx, col, QTableWidgetItem(""))
                self.table_widget.removeCellWidget(row_idx, col)

        self.progress_bar.setVisible(False)
        self.btn_extract.setEnabled(True)
        self.btn_load_meta.setEnabled(True)
        self.btn_load_image.setEnabled(True)
        self.btn_map.setEnabled(True)
        self.statusBar().showMessage("텍스트 추출 완료. 불필요한 항목은 체크 해제 후 매핑을 실행하세요.")

    def on_rect_toggled(self, row_idx, is_included):
        cell_widget = self.table_widget.cellWidget(row_idx, 0)
        if cell_widget:
            checkbox = cell_widget.findChild(QCheckBox)
            if checkbox:
                checkbox.blockSignals(True)
                checkbox.setChecked(is_included)
                checkbox.blockSignals(False)
                
    def on_table_checkbox_toggled(self, row_idx, state):
        is_included = (state == Qt.Checked.value)
        if row_idx < len(self.rect_items):
            rect_item = self.rect_items[row_idx]
            rect_item.is_included = is_included
            rect_item.update_style()
            
    def on_table_selection_changed(self):
        selected_items = self.table_widget.selectedItems()
        if selected_items:
            row = selected_items[0].row()
            if row < len(self.rect_items):
                rect_item = self.rect_items[row]
                # Ensure the selected box is visible in the scroll view
                self.graphics_view.ensureVisible(rect_item, 50, 50)
                # Make it blink
                rect_item.blink()

    def run_map_action(self):
        self.statusBar().showMessage("매핑 중...")
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)
        self.btn_map.setEnabled(False)
        self.btn_extract.setEnabled(False)
        
        texts_to_map = []
        for row in range(self.table_widget.rowCount()):
            cell_widget = self.table_widget.cellWidget(row, 0)
            is_included = cell_widget.findChild(QCheckBox).isChecked() if cell_widget else False
            text = self.table_widget.item(row, 1).text()
            if is_included:
                texts_to_map.append(text)
            else:
                texts_to_map.append(None)
                
        self.worker = WorkerThread(self._map_texts, texts_to_map)
        self.worker.finished.connect(self._on_mapping_finished)
        self.worker.error.connect(self._on_processing_error)
        self.worker.start()

    def _map_texts(self, texts_to_map):
        matches = []
        for text in texts_to_map:
            if text:
                m = self.matcher.find_best_matches([text], top_k=5)[0]
                matches.append(m)
            else:
                matches.append([])
        return matches

    def _on_mapping_finished(self, matches):
        for row_idx, match_list in enumerate(matches):
            cell_widget = self.table_widget.cellWidget(row_idx, 0)
            is_included = cell_widget.findChild(QCheckBox).isChecked() if cell_widget else False
            
            if not is_included or not match_list:
                continue
                
            best_match_kor = match_list[0]['match'] if match_list else ""
            best_match_eng = self.metadata_mapping.get(best_match_kor, "") if best_match_kor else ""
            best_score = match_list[0]['score'] if match_list else 0.0
            
            match_kor_item = QTableWidgetItem(best_match_kor)
            self.table_widget.setItem(row_idx, 2, match_kor_item)
            
            match_eng_item = QTableWidgetItem(best_match_eng)
            self.table_widget.setItem(row_idx, 3, match_eng_item)
            
            score_item = QTableWidgetItem(f"{best_score:.3f}")
            score_item.setFlags(score_item.flags() & ~Qt.ItemIsEditable)
            self.table_widget.setItem(row_idx, 4, score_item)
            
            combo_box = QComboBox()
            combo_box.addItem("--- 선택 ---")
            for m in match_list:
                combo_box.addItem(f"{m['match']} ({m['score']:.2f})", userData=m['match'])
            
            if self.metadata_df is not None:
                combo_box.insertSeparator(combo_box.count())
                for item in self.metadata_mapping.keys():
                    combo_box.addItem(item, userData=item)
            
            combo_box.currentIndexChanged.connect(lambda index, r=row_idx, cb=combo_box: self._update_match_from_combo(r, cb))
            self.table_widget.setCellWidget(row_idx, 5, combo_box)

        self.progress_bar.setVisible(False)
        self.btn_map.setEnabled(True)
        self.btn_extract.setEnabled(True)
        self.btn_save.setEnabled(True)
        self.statusBar().showMessage("매핑 완료.")

    def _update_match_from_combo(self, row, combo_box):
        selected_data = combo_box.currentData() # This is the Korean variable name
        if selected_data:
            eng_data = self.metadata_mapping.get(selected_data, "")
            self.table_widget.setItem(row, 2, QTableWidgetItem(selected_data))
            self.table_widget.setItem(row, 3, QTableWidgetItem(eng_data))
            self.table_widget.setItem(row, 4, QTableWidgetItem("Manual"))

    def _on_processing_error(self, err_msg):
        QMessageBox.critical(self, "오류", f"처리 중 오류 발생: {err_msg}")
        self.progress_bar.setVisible(False)
        self.btn_extract.setEnabled(True)
        self.btn_map.setEnabled(True)
        self.btn_load_meta.setEnabled(True)
        self.btn_load_image.setEnabled(True)
        self.statusBar().showMessage("처리 실패.")

    def save_spec_action(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "명세서 저장", "form_spec.xlsx", "Excel Files (*.xlsx);;JSON Files (*.json);;CSV Files (*.csv)")
        if file_name:
            try:
                final_data = []
                for row in range(self.table_widget.rowCount()):
                    cell_widget = self.table_widget.cellWidget(row, 0)
                    is_included = cell_widget.findChild(QCheckBox).isChecked() if cell_widget else False
                    
                    if not is_included:
                        continue
                        
                    extracted = self.table_widget.item(row, 1).text() if self.table_widget.item(row, 1) else ""
                    mapped_kor = self.table_widget.item(row, 2).text() if self.table_widget.item(row, 2) else ""
                    mapped_eng = self.table_widget.item(row, 3).text() if self.table_widget.item(row, 3) else ""
                    score = self.table_widget.item(row, 4).text() if self.table_widget.item(row, 4) else ""
                    
                    final_data.append({
                        "서식_인식_텍스트": extracted,
                        "추천_변수명(한글)": mapped_kor,
                        "추천_변수명(영문)": mapped_eng,
                        "유사도": score
                    })

                save_spec(final_data, file_name)
                QMessageBox.information(self, "저장 완료", "명세서가 성공적으로 저장되었습니다.")
                
                # Auto-open file
                if sys.platform == "win32":
                    os.startfile(file_name)
                elif sys.platform == "darwin": # macOS
                    subprocess.call(["open", file_name])
                else: # Linux
                    subprocess.call(["xdg-open", file_name])
                    
            except Exception as e:
                QMessageBox.critical(self, "오류", f"저장 중 오류 발생: {e}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
